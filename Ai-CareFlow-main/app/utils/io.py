"""Input/output utilities for AI CareFlow.

This module contains helper functions for exporting and importing
CareFlow artifacts such as summaries or SOAP notes.

At this stage only minimal, file-based helpers are provided without any
external library dependencies.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict
from typing import Any, Dict, List
from io import BytesIO

from fpdf import FPDF
from docx import Document


def save_json_to_file(data: Dict[str, Any], path: str | Path) -> None:
    """Save a dictionary as pretty-printed JSON to ``path``.

    Parameters
    ----------
    data:
        JSON-serializable dictionary to persist.
    path:
        Filesystem path where the JSON file should be written.
    """
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    with target.open("w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def load_json_from_file(path: str | Path) -> Dict[str, Any]:
    """Load JSON data from ``path`` into a dictionary.

    If the file does not exist or cannot be decoded, a ValueError is
    raised to make failures explicit to callers.
    """
    target = Path(path)
    if not target.exists():
        raise ValueError(f"JSON file does not exist: {target}")

    with target.open("r", encoding="utf-8") as f:
        return json.load(f)


def format_soap_note_as_markdown(result: Dict[str, Any]) -> str:
    """Render a SOAP note-style result dictionary as Markdown text.

    Parameters
    ----------
    result:
        A dictionary expected to contain a ``"soap_note"`` mapping with
        ``subjective``, ``objective``, ``assessment``, and ``plan``
        fields, plus an optional ``"summary"`` field.
    """

    soap = result.get("soap_note", {}) or {}
    summary = result.get("summary", "")

    lines: list[str] = []

    if summary:
        lines.append("# Summary")
        if isinstance(summary, list):
            for item in summary:
                lines.append(f"- {item}")
        else:
            lines.append(str(summary))
        lines.append("")

    lines.append("# SOAP Note")
    lines.append("")

    lines.append("## Subjective")
    lines.append(str(soap.get("subjective", "")))
    lines.append("")

    lines.append("## Objective")
    lines.append(str(soap.get("objective", "")))
    lines.append("")

    lines.append("## Assessment")
    lines.append(str(soap.get("assessment", "")))
    lines.append("")

    lines.append("## Plan")
    lines.append(str(soap.get("plan", "")))
    lines.append("")

    return "\n".join(lines)


def to_pretty_json_bytes(data: Dict[str, Any]) -> bytes:
    """Serialize a dictionary to pretty-printed JSON bytes.

    This is convenient for use with ``st.download_button``.
    """

    text = json.dumps(data, indent=2, ensure_ascii=False)
    return text.encode("utf-8")


def build_pdf_from_output(output: Dict[str, Any]) -> bytes:
    """Build a simple PDF representation of the model output.

    The PDF includes Summary, SOAP note, Workflow Suggestions, and
    Missing Information sections, plus a footer disclaimer.
    """

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "AI CareFlow Output", ln=True)

    pdf.ln(4)

    # Summary
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Summary", ln=True)
    pdf.set_font("Arial", "", 11)
    summary = output.get("summary", [])
    if isinstance(summary, list):
        for item in summary:
            pdf.multi_cell(0, 6, f"- {item}")
    else:
        pdf.multi_cell(0, 6, str(summary))

    pdf.ln(3)

    # SOAP Note
    soap = output.get("soap_note", {}) or {}
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "SOAP Note", ln=True)

    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "Subjective", ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(0, 6, str(soap.get("subjective", "")))

    pdf.ln(2)
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "Objective", ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(0, 6, str(soap.get("objective", "")))

    pdf.ln(2)
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "Assessment", ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(0, 6, str(soap.get("assessment", "")))

    pdf.ln(2)
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "Plan", ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(0, 6, str(soap.get("plan", "")))

    pdf.ln(4)

    # Workflow Suggestions
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Workflow Suggestions", ln=True)
    pdf.set_font("Arial", "", 11)
    suggestions: List[str] = output.get("workflow_suggestions", []) or []
    if suggestions:
        for s in suggestions:
            pdf.multi_cell(0, 6, f"- {s}")
    else:
        pdf.multi_cell(0, 6, "No suggestions returned.")

    pdf.ln(4)

    # Missing Information
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Missing Information", ln=True)
    pdf.set_font("Arial", "", 11)
    missing: List[str] = output.get("missing_information", []) or []
    if missing:
        for item in missing:
            pdf.multi_cell(0, 6, f"- {item}")
    else:
        pdf.multi_cell(0, 6, "No clearly missing information identified.")

    pdf.ln(6)

    # Disclaimer
    pdf.set_font("Arial", "I", 9)
    disclaimer = (
        "AI CareFlow is for research and educational purposes only. "
        "Outputs are approximate and must not be used for real patient care."
    )
    pdf.multi_cell(0, 5, disclaimer)

    # fpdf outputs Latin-1 encoded text by default. To avoid
    # UnicodeEncodeError for characters outside this range, fall back
    # to replacing them with a safe placeholder.
    raw = pdf.output(dest="S").encode("latin1", errors="replace")
    return raw


def generate_docx(output: Dict[str, Any]) -> bytes:
    """Build a DOCX file for the SOAP note and related content.

    The document is held in memory and returned as raw bytes.
    """

    doc = Document()

    doc.add_heading("AI CareFlow SOAP Note Draft", level=1)

    summary = output.get("summary", [])
    soap = output.get("soap_note", {}) or {}
    suggestions: List[str] = output.get("workflow_suggestions", []) or []

    # Summary
    doc.add_heading("Summary", level=2)
    if isinstance(summary, list):
        for item in summary:
            doc.add_paragraph(str(item), style="List Bullet")
    elif summary:
        doc.add_paragraph(str(summary))

    # SOAP sections
    doc.add_heading("Subjective", level=2)
    doc.add_paragraph(str(soap.get("subjective", "")))

    doc.add_heading("Objective", level=2)
    doc.add_paragraph(str(soap.get("objective", "")))

    doc.add_heading("Assessment", level=2)
    doc.add_paragraph(str(soap.get("assessment", "")))

    doc.add_heading("Plan", level=2)
    doc.add_paragraph(str(soap.get("plan", "")))

    # Workflow Suggestions
    doc.add_heading("Workflow Suggestions", level=2)
    if suggestions:
        for s in suggestions:
            doc.add_paragraph(str(s), style="List Bullet")
    else:
        doc.add_paragraph("No suggestions returned.")

    # Footer disclaimer
    doc.add_paragraph(
        "AI CareFlow is for research and educational purposes only. "
        "Outputs are approximate and must not be used for real patient care.",
        style=None,
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


